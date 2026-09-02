from __future__ import annotations

import csv
import io
from pathlib import Path
from typing import BinaryIO
from urllib.parse import urlparse
from urllib.request import Request, urlopen

from bs4 import BeautifulSoup
from django.db import transaction
from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from apps.ai_engagement.models import (
    Chunk,
    Document,
    KnowledgeSource,
)


class KnowledgeExtractionError(Exception):
    """
    Raised when SHVYA cannot extract readable text from a
    supported knowledge source.
    """


class KnowledgeIngestionService:
    """
    Extract, clean, chunk, and persist organization knowledge.

    Supported file types:
        .txt
        .csv
        .xlsx
        .pdf
        .docx

    Supported URL sources:
        http://
        https://

    Main pipeline:

        Source
            ↓
        Extraction
            ↓
        Cleaning
            ↓
        Versioning
            ↓
        Chunking
            ↓
        Chunk records

    This service intentionally does NOT:
        - call an LLM
        - generate embeddings
        - perform vector search
        - modify leads
        - send messages

    Embeddings and retrieval are separate layers.
    """

    # ============================================================
    # SUPPORTED FILE TYPES
    # ============================================================

    SUPPORTED_FILE_EXTENSIONS = {
        ".txt",
        ".csv",
        ".xlsx",
        ".pdf",
        ".docx",
    }

    # ============================================================
    # CONFIGURATION
    # ============================================================

    REQUEST_TIMEOUT_SECONDS = 20

    DEFAULT_CHUNK_SIZE = 1800

    DEFAULT_CHUNK_OVERLAP = 200

    USER_AGENT = (
        "SHVYA-Knowledge-Bot/1.0 "
        "(knowledge ingestion)"
    )

    # ============================================================
    # PUBLIC: DOCUMENT INGESTION
    # ============================================================

    def ingest_document(
        self,
        document: Document,
    ) -> int:
        """
        Process an uploaded Document.

        The supplied Document represents the new version of a
        logical knowledge source.

        Flow:

            Document
                ↓
            Validate source
                ↓
            Extract text
                ↓
            Clean text
                ↓
            Chunk
                ↓
            Determine version
                ↓
            Deactivate previous version
                ↓
            Save chunks
                ↓
            Mark COMPLETED

        Returns:
            Number of chunks created.

        Raises:
            KnowledgeExtractionError
        """

        document.processing_status = (
            Document.ProcessingStatus.PROCESSING
        )

        document.processing_error = ""

        document.save(
            update_fields=[
                "processing_status",
                "processing_error",
                "updated_at",
            ]
        )

        try:
            if not document.file:
                raise KnowledgeExtractionError(
                    "Document does not contain an uploaded file."
                )

            filename = Path(
                document.file.name
            ).name

            # ----------------------------------------------------
            # Determine stable logical source identity.
            # ----------------------------------------------------

            source_key = (
                document.source_key
                or document.name
                or filename
            ).strip()

            if not source_key:
                raise KnowledgeExtractionError(
                    "Document source_key cannot be empty."
                )

            # ----------------------------------------------------
            # Extract source text.
            # ----------------------------------------------------

            document.file.seek(0)

            text = self.extract_file_text(
                document.file,
                filename=filename,
            )

            if not text.strip():
                raise KnowledgeExtractionError(
                    "No readable text was extracted from the document."
                )

            # ----------------------------------------------------
            # Create chunks.
            # ----------------------------------------------------

            chunks = self.chunk_text(
                text
            )

            if not chunks:
                raise KnowledgeExtractionError(
                    "No usable chunks were created from the document."
                )

            with transaction.atomic():

                # ------------------------------------------------
                # Lock existing versions for this source.
                # ------------------------------------------------

                latest = (
                    Document.objects
                    .select_for_update()
                    .filter(
                        organization=document.organization,
                        source_key=source_key,
                    )
                    .exclude(
                        pk=document.pk,
                    )
                    .order_by(
                        "-version",
                    )
                    .first()
                )

                next_version = (
                    latest.version + 1
                    if latest
                    else 1
                )

                # ------------------------------------------------
                # Deactivate previous active version.
                # ------------------------------------------------

                Document.objects.filter(
                    organization=document.organization,
                    source_key=source_key,
                    is_active=True,
                ).exclude(
                    pk=document.pk,
                ).update(
                    is_active=False,
                )

                # ------------------------------------------------
                # Make the supplied document the new version.
                # ------------------------------------------------

                document.source_key = source_key

                document.version = next_version

                document.is_active = True

                document.processing_status = (
                    Document.ProcessingStatus.COMPLETED
                )

                document.processing_error = ""

                document.save(
                    update_fields=[
                        "source_key",
                        "version",
                        "is_active",
                        "processing_status",
                        "processing_error",
                        "updated_at",
                    ]
                )

                # ------------------------------------------------
                # Ensure this version starts with no old chunks.
                # ------------------------------------------------

                document.chunks.all().delete()

                # ------------------------------------------------
                # Persist chunks.
                # Embeddings will be added by the embedding layer.
                # ------------------------------------------------

                Chunk.objects.bulk_create(
                    [
                        Chunk(
                            document=document,
                            organization=document.organization,
                            content=chunk,
                            chunk_index=index,
                            is_active=True,
                        )
                        for index, chunk in enumerate(
                            chunks
                        )
                    ]
                )

            return len(chunks)

        except Exception as exc:

            # ----------------------------------------------------
            # Preserve failure state outside the successful
            # transaction so FAILED is not rolled back.
            # ----------------------------------------------------

            document.processing_status = (
                Document.ProcessingStatus.FAILED
            )

            document.processing_error = str(
                exc
            )

            document.save(
                update_fields=[
                    "processing_status",
                    "processing_error",
                    "updated_at",
                ]
            )

            if isinstance(
                exc,
                KnowledgeExtractionError,
            ):
                raise

            raise KnowledgeExtractionError(
                f"Unable to ingest document "
                f"{document.name}: {exc}"
            ) from exc

        # ============================================================
    # VERSIONED URL INGESTION
    # ============================================================

    def ingest_url(
        self,
        source: KnowledgeSource,
    ) -> int:
        """
        Fetch, clean, chunk, and persist a URL as a new
        Document version.

        Publication rules:

            1. Existing active version remains active while the
               new source is being fetched and processed.

            2. New version is created as PROCESSING.

            3. Chunks are created.

            4. New version is marked COMPLETED.

            5. Previous active versions are deactivated.

            6. New version becomes ACTIVE.

        If anything fails before publication, the previous active
        version remains available.
        """

        if source.source_type != (
            KnowledgeSource.SourceType.URL
        ):
            raise KnowledgeExtractionError(
                "KnowledgeSource is not a URL source."
            )

        if not source.url:
            raise KnowledgeExtractionError(
                "KnowledgeSource does not contain a URL."
            )

        document = None
        document_created = False

        try:

            # ----------------------------------------------------
            # Normalize source URL.
            # ----------------------------------------------------

            normalized_url = (
                self._normalize_url(
                    source.url
                )
            )

            source_key = normalized_url

            document_name = (
                source.name.strip()
                if source.name
                else normalized_url
            )

            # ----------------------------------------------------
            # Fetch and process the source BEFORE changing
            # the active published version.
            # ----------------------------------------------------

            text = self.extract_url_text(
                normalized_url
            )

            if not text.strip():
                raise KnowledgeExtractionError(
                    "No readable text was extracted from the URL."
                )

            chunks = self.chunk_text(
                text
            )

            if not chunks:
                raise KnowledgeExtractionError(
                    "No usable chunks were created from the URL."
                )

            # ----------------------------------------------------
            # Determine the next version.
            # Lock existing versions only while calculating
            # version + publishing the new version.
            # ----------------------------------------------------

            with transaction.atomic():

                existing_versions = list(
                    Document.objects
                    .select_for_update()
                    .filter(
                        organization=source.organization,
                        source_key=source_key,
                    )
                    .order_by(
                        "-version",
                    )
                )

                latest = (
                    existing_versions[0]
                    if existing_versions
                    else None
                )

                next_version = (
                    latest.version + 1
                    if latest
                    else 1
                )

                # ------------------------------------------------
                # Create the new version as PROCESSING.
                #
                # The previous active version remains active here.
                # ------------------------------------------------

                document = Document.objects.create(
                    organization=source.organization,
                    name=document_name,
                    source_key=source_key,
                    version=next_version,
                    source_url=normalized_url,
                    processing_status=(
                        Document.ProcessingStatus.PROCESSING
                    ),
                    processing_error="",
                    is_active=False,
                )

                document_created = True

                # ------------------------------------------------
                # Create chunks while the new version is still
                # unpublished.
                # ------------------------------------------------

                Chunk.objects.bulk_create(
                    [
                        Chunk(
                            document=document,
                            organization=source.organization,
                            content=chunk,
                            chunk_index=index,
                            is_active=True,
                        )
                        for index, chunk in enumerate(
                            chunks
                        )
                    ]
                )

                # ------------------------------------------------
                # The new document is now fully processed.
                # ------------------------------------------------

                document.processing_status = (
                    Document.ProcessingStatus.COMPLETED
                )

                document.processing_error = ""

                document.save(
                    update_fields=[
                        "processing_status",
                        "processing_error",
                        "updated_at",
                    ]
                )

                # ------------------------------------------------
                # Only after successful processing do we publish
                # the new version.
                # ------------------------------------------------

                Document.objects.filter(
                    organization=source.organization,
                    source_key=source_key,
                    is_active=True,
                ).exclude(
                    pk=document.pk,
                ).update(
                    is_active=False,
                )

                document.is_active = True

                document.save(
                    update_fields=[
                        "is_active",
                        "updated_at",
                    ]
                )

            return len(chunks)

        except Exception as exc:

            # ----------------------------------------------------
            # If the database transaction rolled back, the newly
            # created Document may no longer exist.
            #
            # Therefore only update it when the row still exists.
            # ----------------------------------------------------

            if (
                document is not None
                and document_created
                and document.pk
                and Document.objects.filter(
                    pk=document.pk,
                ).exists()
            ):

                document.processing_status = (
                    Document.ProcessingStatus.FAILED
                )

                document.processing_error = str(
                    exc
                )

                document.is_active = False

                document.save(
                    update_fields=[
                        "processing_status",
                        "processing_error",
                        "is_active",
                        "updated_at",
                    ]
                )

            if isinstance(
                exc,
                KnowledgeExtractionError,
            ):
                raise

            raise KnowledgeExtractionError(
                f"Unable to ingest URL "
                f"{source.url}: {exc}"
            ) from exc

    # ============================================================
    # FILE EXTRACTION
    # ============================================================

    def extract_file_text(
        self,
        file_obj: BinaryIO,
        *,
        filename: str,
    ) -> str:
        """
        Extract readable text from a supported file.
        """

        extension = (
            Path(filename)
            .suffix
            .lower()
            .strip()
        )

        if not extension:
            raise KnowledgeExtractionError(
                "The uploaded file has no extension."
            )

        if extension not in self.SUPPORTED_FILE_EXTENSIONS:
            raise KnowledgeExtractionError(
                f"Unsupported file type: {extension}"
            )

        try:

            if extension == ".txt":
                return self._extract_txt(
                    file_obj
                )

            if extension == ".csv":
                return self._extract_csv(
                    file_obj
                )

            if extension == ".xlsx":
                return self._extract_excel(
                    file_obj
                )

            if extension == ".pdf":
                return self._extract_pdf(
                    file_obj
                )

            if extension == ".docx":
                return self._extract_docx(
                    file_obj
                )

        except KnowledgeExtractionError:
            raise

        except Exception as exc:
            raise KnowledgeExtractionError(
                f"Unable to extract text from "
                f"{filename}: {exc}"
            ) from exc

        raise KnowledgeExtractionError(
            f"No extractor available for {extension}"
        )

    # ============================================================
    # TXT
    # ============================================================

    def _extract_txt(
        self,
        file_obj: BinaryIO,
    ) -> str:
        """
        Extract UTF-8 text from a plain-text file.
        """

        raw = file_obj.read()

        if isinstance(raw, str):
            text = raw
        else:
            text = raw.decode(
                "utf-8",
                errors="replace",
            )

        return self._clean_text(
            text
        )

    # ============================================================
    # CSV
    # ============================================================

    def _extract_csv(
        self,
        file_obj: BinaryIO,
    ) -> str:
        """
        Convert CSV rows into readable text.
        """

        raw = file_obj.read()

        if isinstance(raw, str):
            text = raw
        else:
            text = raw.decode(
                "utf-8-sig",
                errors="replace",
            )

        stream = io.StringIO(
            text
        )

        reader = csv.reader(
            stream
        )

        rows = []

        for row in reader:

            values = [
                str(value).strip()
                for value in row
                if str(value).strip()
            ]

            if values:
                rows.append(
                    " | ".join(values)
                )

        return self._clean_text(
            "\n".join(rows)
        )

    # ============================================================
    # XLSX
    # ============================================================

    def _extract_excel(
        self,
        file_obj: BinaryIO,
    ) -> str:
        """
        Extract readable cell values from XLSX workbooks.
        """

        workbook = load_workbook(
            filename=file_obj,
            read_only=True,
            data_only=True,
        )

        rows = []

        try:

            for worksheet in workbook.worksheets:

                rows.append(
                    f"Sheet: {worksheet.title}"
                )

                for row in worksheet.iter_rows(
                    values_only=True
                ):

                    values = []

                    for value in row:

                        if value is None:
                            continue

                        value_text = str(
                            value
                        ).strip()

                        if value_text:
                            values.append(
                                value_text
                            )

                    if values:
                        rows.append(
                            " | ".join(values)
                        )

        finally:

            workbook.close()

        return self._clean_text(
            "\n".join(rows)
        )

    # ============================================================
    # PDF
    # ============================================================

    def _extract_pdf(
        self,
        file_obj: BinaryIO,
    ) -> str:
        """
        Extract readable text from PDF pages.
        """

        reader = PdfReader(
            file_obj
        )

        pages = []

        for index, page in enumerate(
            reader.pages,
            start=1,
        ):

            text = page.extract_text()

            if not text:
                continue

            text = text.strip()

            if not text:
                continue

            pages.append(
                f"Page {index}\n{text}"
            )

        return self._clean_text(
            "\n\n".join(pages)
        )

    # ============================================================
    # DOCX
    # ============================================================

    def _extract_docx(
        self,
        file_obj: BinaryIO,
    ) -> str:
        """
        Extract paragraphs and tables from a DOCX file.
        """

        document = DocxDocument(
            file_obj
        )

        paragraphs = []

        for paragraph in document.paragraphs:

            text = (
                paragraph.text or ""
            ).strip()

            if text:
                paragraphs.append(
                    text
                )

        table_rows = []

        for table in document.tables:

            for row in table.rows:

                values = []

                for cell in row.cells:

                    text = (
                        cell.text or ""
                    ).strip()

                    if text:
                        values.append(
                            text
                        )

                if values:
                    table_rows.append(
                        " | ".join(values)
                    )

        combined = []

        if paragraphs:
            combined.extend(
                paragraphs
            )

        if table_rows:

            if combined:
                combined.append("")

            combined.extend(
                table_rows
            )

        return self._clean_text(
            "\n".join(combined)
        )

    # ============================================================
    # URL EXTRACTION
    # ============================================================

    def extract_url_text(
        self,
        url: str,
    ) -> str:
        """
        Fetch and extract readable text from an HTTP/HTTPS page.
        """

        normalized_url = (
            self._normalize_url(
                url
            )
        )

        request = Request(
            normalized_url,
            headers={
                "User-Agent":
                    self.USER_AGENT,
            },
        )

        try:

            with urlopen(
                request,
                timeout=self.REQUEST_TIMEOUT_SECONDS,
            ) as response:

                content_type = (
                    response.headers.get(
                        "Content-Type",
                        "",
                    )
                    .lower()
                )

                raw = response.read()

        except Exception as exc:

            raise KnowledgeExtractionError(
                f"Unable to fetch URL "
                f"{normalized_url}: {exc}"
            ) from exc

        if (
            "text/html" not in content_type
            and "application/xhtml+xml"
            not in content_type
        ):
            raise KnowledgeExtractionError(
                "The URL did not return an HTML page."
            )

        try:

            html = raw.decode(
                "utf-8",
                errors="replace",
            )

            soup = BeautifulSoup(
                html,
                "html.parser",
            )

            self._remove_unwanted_html(
                soup
            )

            text = soup.get_text(
                separator="\n"
            )

            return self._clean_text(
                text
            )

        except Exception as exc:

            raise KnowledgeExtractionError(
                f"Unable to parse URL "
                f"{normalized_url}: {exc}"
            ) from exc

    # ============================================================
    # URL HELPERS
    # ============================================================

    def normalize_url(
        self,
        url: str,
    ) -> str:
        """
        Public wrapper around URL normalization.

        ingest_url() publishes the new Document under
        source_key == this normalized URL but does not return the
        Document itself. Callers (e.g. the ingestion Celery task)
        use this method to resolve that same source_key afterwards
        instead of reaching into the private normalizer.
        """

        return self._normalize_url(
            url
        )

    def _normalize_url(
        self,
        url: str,
    ) -> str:
        """
        Normalize and validate HTTP/HTTPS URLs.
        """

        normalized = (
            url or ""
        ).strip()

        if not normalized:
            raise KnowledgeExtractionError(
                "URL cannot be empty."
            )

        parsed = urlparse(
            normalized
        )

        if not parsed.scheme:

            normalized = (
                "https://"
                + normalized
            )

            parsed = urlparse(
                normalized
            )

        if parsed.scheme not in {
            "http",
            "https",
        }:
            raise KnowledgeExtractionError(
                "Only HTTP and HTTPS URLs "
                "are supported."
            )

        if not parsed.netloc:
            raise KnowledgeExtractionError(
                "Invalid URL."
            )

        return normalized

    def _remove_unwanted_html(
        self,
        soup: BeautifulSoup,
    ) -> None:
        """
        Remove HTML elements that do not contain useful
        knowledge content.
        """

        for tag in soup(
            [
                "script",
                "style",
                "noscript",
                "svg",
                "canvas",
                "template",
            ]
        ):
            tag.decompose()

    # ============================================================
    # TEXT CLEANING
    # ============================================================

    def _clean_text(
        self,
        text: str,
    ) -> str:
        """
        Normalize whitespace while preserving useful line
        boundaries for chunking.
        """

        if not text:
            return ""

        lines = []

        for raw_line in text.splitlines():

            line = (
                " ".join(
                    raw_line.split()
                )
                .strip()
            )

            if line:
                lines.append(
                    line
                )

        cleaned = []

        previous = None

        for line in lines:

            if line == previous:
                continue

            cleaned.append(
                line
            )

            previous = line

        return "\n".join(
            cleaned
        )

    # ============================================================
    # TEXT CHUNKING
    # ============================================================

    def chunk_text(
        self,
        text: str,
        *,
        max_chars: int = DEFAULT_CHUNK_SIZE,
        overlap: int = DEFAULT_CHUNK_OVERLAP,
    ) -> list[str]:
        """
        Split cleaned text into overlapping chunks.

        The chunker prefers line boundaries and falls back to
        character splitting for a single oversized line.
        """

        if not text:
            return []

        if max_chars <= 0:
            raise ValueError(
                "max_chars must be greater than zero."
            )

        if overlap < 0:
            raise ValueError(
                "overlap cannot be negative."
            )

        if overlap >= max_chars:
            raise ValueError(
                "overlap must be smaller than max_chars."
            )

        lines = [
            line.strip()
            for line in text.splitlines()
            if line.strip()
        ]

        if not lines:
            return []

        chunks: list[str] = []

        current_lines: list[str] = []

        current_length = 0

        for line in lines:

            line_length = len(line)

            # ----------------------------------------------------
            # Oversized line.
            # ----------------------------------------------------

            if line_length > max_chars:

                if current_lines:

                    chunks.append(
                        "\n".join(
                            current_lines
                        )
                    )

                    current_lines = []

                    current_length = 0

                start = 0

                while start < line_length:

                    end = min(
                        start + max_chars,
                        line_length,
                    )

                    piece = (
                        line[start:end]
                        .strip()
                    )

                    if piece:
                        chunks.append(
                            piece
                        )

                    if end >= line_length:
                        break

                    next_start = (
                        end - overlap
                    )

                    if next_start <= start:
                        next_start = (
                            start + 1
                        )

                    start = next_start

                continue

            # ----------------------------------------------------
            # Normal line.
            # ----------------------------------------------------

            separator_length = (
                1
                if current_lines
                else 0
            )

            proposed_length = (
                current_length
                + separator_length
                + line_length
            )

            # ----------------------------------------------------
            # Flush current chunk if necessary.
            # ----------------------------------------------------

            if (
                current_lines
                and proposed_length > max_chars
            ):

                chunks.append(
                    "\n".join(
                        current_lines
                    )
                )

                # ------------------------------------------------
                # Build overlap from previous lines.
                # ------------------------------------------------

                overlap_lines: list[str] = []

                overlap_length = 0

                for previous_line in reversed(
                    current_lines
                ):

                    additional_length = (
                        len(previous_line)
                        + (
                            1
                            if overlap_lines
                            else 0
                        )
                    )

                    if (
                        overlap_length
                        + additional_length
                        > overlap
                    ):
                        break

                    overlap_lines.insert(
                        0,
                        previous_line,
                    )

                    overlap_length += (
                        additional_length
                    )

                current_lines = (
                    overlap_lines
                )

                current_length = (
                    sum(
                        len(item)
                        for item in current_lines
                    )
                    + max(
                        len(current_lines) - 1,
                        0,
                    )
                )

            # ----------------------------------------------------
            # Add line to current chunk.
            # ----------------------------------------------------

            current_lines.append(
                line
            )

            current_length = (
                sum(
                    len(item)
                    for item in current_lines
                )
                + max(
                    len(current_lines) - 1,
                    0,
                )
            )

        # --------------------------------------------------------
        # Flush final chunk.
        # --------------------------------------------------------

        if current_lines:

            chunks.append(
                "\n".join(
                    current_lines
                )
            )

        return [
            chunk.strip()
            for chunk in chunks
            if chunk.strip()
        ]